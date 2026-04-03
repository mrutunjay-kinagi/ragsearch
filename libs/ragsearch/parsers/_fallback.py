"""
Pure-Python fallback document parser.
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterator

from ..errors import ParseCorruptError, ParserUnavailableError, UnsupportedFileTypeError
from ._models import ParsedDocument


class FallbackParser:
    """Pure Python parser used when LiteParse is unavailable."""

    SUPPORTED_SUFFIXES = {".txt", ".md", ".html", ".htm", ".pdf", ".docx", ".doc"}

    def supports(self, path: Path | str) -> bool:
        path = Path(path) if not isinstance(path, Path) else path
        return path.suffix.lower() in self.SUPPORTED_SUFFIXES

    def parse(self, path: Path | str) -> Iterator[ParsedDocument]:
        if path is None:
            raise ParseCorruptError("Input path cannot be None")
        if not isinstance(path, Path):
            path = Path(path)
        if not path.exists():
            raise ParseCorruptError(f"Input path is missing or unreadable: {path}")
        if not self.supports(path):
            raise UnsupportedFileTypeError(f"Unsupported file type: {path.suffix} for {path}")

        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"}:
            yield ParsedDocument(
                text=path.read_text(encoding="utf-8"),
                metadata={"source_path": str(path), "parser_name": "fallback/plain_text"},
                source_path=str(path),
                parser_name="fallback/plain_text",
            )
            return

        if suffix in {".html", ".htm"}:
            try:
                from bs4 import BeautifulSoup
            except ImportError as exc:
                raise ParserUnavailableError("Install ragsearch[parsers] to parse HTML files", cause=exc) from exc

            text = BeautifulSoup(path.read_text(encoding="utf-8"), "html.parser").get_text(" ", strip=True)
            yield ParsedDocument(
                text=text,
                metadata={"source_path": str(path), "parser_name": "fallback/bs4"},
                source_path=str(path),
                parser_name="fallback/bs4",
            )
            return

        if suffix == ".pdf":
            try:
                from pypdf import PdfReader
            except ImportError as exc:
                raise ParserUnavailableError("Install ragsearch[parsers] to parse PDF files", cause=exc) from exc

            try:
                reader = PdfReader(str(path))
                text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
            except Exception as exc:
                raise ParseCorruptError(f"Failed to parse PDF file: {path}", cause=exc) from exc

            yield ParsedDocument(
                text=text,
                metadata={"source_path": str(path), "parser_name": "fallback/pypdf"},
                source_path=str(path),
                parser_name="fallback/pypdf",
            )
            return

        if suffix in {".docx", ".doc"}:
            try:
                import docx
            except ImportError as exc:
                raise ParserUnavailableError("Install ragsearch[parsers] to parse DOCX/DOC files", cause=exc) from exc

            try:
                document = docx.Document(str(path))
                text = "\n".join(paragraph.text for paragraph in document.paragraphs).strip()
            except Exception as exc:
                raise ParseCorruptError(f"Failed to parse DOCX/DOC file: {path}", cause=exc) from exc

            yield ParsedDocument(
                text=text,
                metadata={"source_path": str(path), "parser_name": "fallback/python-docx"},
                source_path=str(path),
                parser_name="fallback/python-docx",
            )
            return

        raise UnsupportedFileTypeError(f"Unsupported file type: {path.suffix} for {path}")
